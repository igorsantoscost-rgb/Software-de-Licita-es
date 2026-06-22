"""
Extracao de texto de documentos anexados a licitacao (edital, termo de
referencia, outros), para alimentar o Resumo da Oportunidade gerado por IA.

PDFs nao sao extraidos aqui - sao enviados diretamente como documento nativo
para a API da Anthropic, que le PDF sem precisar de OCR/extracao no servidor.
"""
import os
import csv
import io


def extrair_texto_docx(caminho):
    from docx import Document
    doc = Document(caminho)
    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas = [c.text.strip() for c in linha.cells]
            if any(celulas):
                partes.append(" | ".join(celulas))
    return "\n".join(partes)


def extrair_texto_xlsx(caminho):
    from openpyxl import load_workbook
    wb = load_workbook(caminho, data_only=True, read_only=True)
    partes = []
    for aba in wb.worksheets:
        partes.append(f"--- Planilha: {aba.title} ---")
        for linha in aba.iter_rows(values_only=True):
            valores = [str(v) for v in linha if v is not None]
            if valores:
                partes.append(" | ".join(valores))
    return "\n".join(partes)


def extrair_texto_csv(caminho):
    partes = []
    with open(caminho, newline="", encoding="utf-8", errors="ignore") as f:
        leitor = csv.reader(f)
        for linha in leitor:
            if any(c.strip() for c in linha):
                partes.append(" | ".join(linha))
    return "\n".join(partes)


def extrair_texto_html(caminho):
    from bs4 import BeautifulSoup
    with open(caminho, encoding="utf-8", errors="ignore") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def extrair_texto_txt(caminho):
    with open(caminho, encoding="utf-8", errors="ignore") as f:
        return f.read()


# Tamanho maximo de texto extraido por documento, para nao explodir o prompt
LIMITE_CHARS_POR_DOC = 40000


def extrair_texto_documento(doc):
    """Extrai texto de um Documento (modelo) que NAO seja PDF.
    Retorna string vazia se o formato nao for suportado ou o arquivo nao existir.
    """
    if not os.path.exists(doc.caminho):
        return ""

    ext = os.path.splitext(doc.nome_original)[1].lower()
    try:
        if ext == ".docx":
            texto = extrair_texto_docx(doc.caminho)
        elif ext in (".xlsx", ".xlsm"):
            texto = extrair_texto_xlsx(doc.caminho)
        elif ext == ".csv":
            texto = extrair_texto_csv(doc.caminho)
        elif ext in (".html", ".htm"):
            texto = extrair_texto_html(doc.caminho)
        elif ext == ".txt":
            texto = extrair_texto_txt(doc.caminho)
        else:
            return ""
    except Exception as e:
        return f"[Nao foi possivel ler o arquivo {doc.nome_original}: {e}]"

    if len(texto) > LIMITE_CHARS_POR_DOC:
        texto = texto[:LIMITE_CHARS_POR_DOC] + "\n[...texto truncado...]"
    return texto


def ler_pdf_base64(doc):
    """Le um Documento PDF e retorna seus bytes em base64, para envio nativo
    a API da Anthropic. Retorna None se o arquivo nao existir."""
    import base64
    if not os.path.exists(doc.caminho):
        return None
    with open(doc.caminho, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")
